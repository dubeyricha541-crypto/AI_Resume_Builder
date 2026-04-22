import os

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import get_object_or_404

from .models import Resume


@login_required
def download_resume_docx_unified(request, pk):
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt, RGBColor
    except Exception as e:
        return HttpResponse(f"python-docx is not installed: {e}", status=500)

    def set_cell_border(cell, **kwargs):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)

        for edge in ("left", "top", "right", "bottom"):
            edge_data = kwargs.get(edge)
            if not edge_data:
                continue
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))

    def hide_cell_borders(cell):
        set_cell_border(
            cell,
            top={"val": "nil"},
            bottom={"val": "nil"},
            left={"val": "nil"},
            right={"val": "nil"},
        )

    def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.line_spacing = line

    def style_run(run, size, *, bold=False, color="111111"):
        run.bold = bold
        run.font.name = "Segoe UI"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)

    def add_section_heading(doc, title):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell = table.cell(0, 0)
        cell.width = Mm(178)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(
            cell,
            left={"val": "nil"},
            right={"val": "nil"},
            top={"val": "nil"},
            bottom={"val": "single", "sz": 8, "space": 0, "color": "000000"},
        )
        paragraph = cell.paragraphs[0]
        set_paragraph_spacing(paragraph, before=0, after=2, line=1.0)
        run = paragraph.add_run(title.upper())
        style_run(run, 10, bold=True)

    def add_body_text(doc, text):
        paragraph = doc.add_paragraph()
        set_paragraph_spacing(paragraph, before=0, after=3, line=1.08)
        run = paragraph.add_run(text)
        style_run(run, 9.5)

    def add_two_col_entry(doc, primary, secondary="", dates="", description=""):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        left = table.rows[0].cells[0]
        right = table.rows[0].cells[1]
        left.width = Mm(140)
        right.width = Mm(38)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        hide_cell_borders(left)
        hide_cell_borders(right)

        title_p = left.paragraphs[0]
        set_paragraph_spacing(title_p, before=0, after=0, line=1.0)
        title_run = title_p.add_run(primary)
        style_run(title_run, 10, bold=True)

        if secondary:
            secondary_p = left.add_paragraph()
            set_paragraph_spacing(secondary_p, before=0, after=1, line=1.0)
            secondary_run = secondary_p.add_run(secondary)
            style_run(secondary_run, 10)
            secondary_run.bold = True

        if description:
            desc_p = left.add_paragraph()
            set_paragraph_spacing(desc_p, before=0, after=0, line=1.08)
            desc_run = desc_p.add_run(description)
            style_run(desc_run, 9.5)

        date_p = right.paragraphs[0]
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(date_p, before=0, after=0, line=1.0)
        date_run = date_p.add_run(dates)
        style_run(date_run, 9, color="333333")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Segoe UI"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
    normal_style.font.size = Pt(9.5)

    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False

    left = header.rows[0].cells[0]
    right = header.rows[0].cells[1]
    left.width = Mm(140)
    right.width = Mm(38)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_border(
        left,
        left={"val": "nil"},
        top={"val": "nil"},
        right={"val": "nil"},
        bottom={"val": "single", "sz": 8, "space": 0, "color": "000000"},
    )
    set_cell_border(
        right,
        left={"val": "nil"},
        top={"val": "nil"},
        right={"val": "nil"},
        bottom={"val": "single", "sz": 8, "space": 0, "color": "000000"},
    )

    name_p = left.paragraphs[0]
    set_paragraph_spacing(name_p, before=0, after=0, line=1.0)
    name_run = name_p.add_run(resume.full_name or "Resume")
    style_run(name_run, 24, bold=True)

    if resume.job_title:
        title_p = left.add_paragraph()
        set_paragraph_spacing(title_p, before=0, after=2, line=1.0)
        title_run = title_p.add_run(resume.job_title)
        style_run(title_run, 11, bold=True)

    contact_lines = [value for value in [resume.email, resume.phone, resume.address] if value]
    if contact_lines:
        contact_p = left.add_paragraph()
        set_paragraph_spacing(contact_p, before=0, after=0, line=1.0)
        for index, line in enumerate(contact_lines):
            run = contact_p.add_run(line)
            style_run(run, 9, color="333333")
            if index < len(contact_lines) - 1:
                run.add_break()

    if resume.photo and getattr(resume.photo, "path", None) and os.path.exists(resume.photo.path):
        photo_p = right.paragraphs[0]
        photo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(photo_p, before=0, after=0, line=1.0)
        photo_p.add_run().add_picture(resume.photo.path, width=Mm(24), height=Mm(24))

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=5, line=1.0)

    if resume.career_objective:
        add_section_heading(doc, "Summary")
        add_body_text(doc, resume.career_objective)

    if resume.skills:
        add_section_heading(doc, "Skills")
        skills_p = doc.add_paragraph()
        set_paragraph_spacing(skills_p, before=0, after=4, line=1.08)
        skills_text = ", ".join([s.strip() for s in resume.skills.split(",") if s.strip()])
        skills_run = skills_p.add_run(skills_text or resume.skills)
        style_run(skills_run, 9.5)

    exp_entries = resume.exp_entries() if hasattr(resume, "exp_entries") else []
    if resume.experience:
        add_section_heading(doc, "Experience")
        if exp_entries:
            for entry in exp_entries:
                add_two_col_entry(
                    doc,
                    entry.get("title", ""),
                    entry.get("company", ""),
                    " - ".join([value for value in [entry.get("start"), entry.get("end")] if value]),
                    entry.get("desc", ""),
                )
        else:
            add_body_text(doc, resume.experience)

    edu_entries = resume.edu_entries() if hasattr(resume, "edu_entries") else []
    if resume.education:
        add_section_heading(doc, "Education")
        if edu_entries:
            for entry in edu_entries:
                add_two_col_entry(
                    doc,
                    entry.get("degree", ""),
                    entry.get("institution", ""),
                    " - ".join([value for value in [entry.get("start"), entry.get("end")] if value]),
                )
        else:
            add_body_text(doc, resume.education)

    if resume.projects:
        add_section_heading(doc, "Projects")
        add_body_text(doc, resume.projects)
    if resume.certifications:
        add_section_heading(doc, "Certifications")
        add_body_text(doc, resume.certifications)
    if resume.achievements:
        add_section_heading(doc, "Achievements")
        add_body_text(doc, resume.achievements)
    if resume.languages:
        add_section_heading(doc, "Languages")
        add_body_text(doc, resume.languages)

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="resume_{resume.id}.docx"'
    doc.save(response)
    return response
