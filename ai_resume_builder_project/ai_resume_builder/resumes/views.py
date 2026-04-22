from io import BytesIO

from django.views.generic import ListView, UpdateView, DeleteView, DetailView
from django.urls import reverse_lazy, reverse
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.decorators import login_required
from django.template.loader import get_template
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, get_object_or_404, redirect
from django.views.decorators.http import require_http_methods
from xhtml2pdf import pisa
import base64
import uuid
import os
from django.core.files.base import ContentFile
from django.conf import settings
import json

from .models import Resume
from .forms import ResumeForm
from .ai_utils import generate_resume, build_improvement_prompt, build_ats_prompt
from .ats_scoring import (
    compute_ats_score,
    heuristic_ats_score,
    refresh_resume_ats_score,
)


ATS_TEMPLATE = Resume.ATS_TEMPLATE


# ✅ CREATE VIEW — unified form with live preview (No login required to create)
def create_resume(request):
    # Check if user is logged in to determine redirect behavior
    is_authenticated = request.user.is_authenticated
    
    if request.method == 'POST':
        # Check if user is logged in when saving
        if not is_authenticated:
            return redirect(f"{settings.LOGIN_URL}?next={request.path}")
        
        # Collect base fields
        resume = Resume(
            user=request.user,
            template=ATS_TEMPLATE,
            full_name=request.POST.get('full_name', '').strip(),
            job_title=request.POST.get('job_title', '').strip(),
            email=request.POST.get('email', '').strip(),
            phone=request.POST.get('phone', '').strip(),
            address=request.POST.get('address', '').strip(),
            linkedin=request.POST.get('linkedin', '').strip(),
            portfolio=request.POST.get('portfolio', '').strip(),
            github=request.POST.get('github', '').strip(),
            twitter=request.POST.get('twitter', '').strip(),
            career_objective=request.POST.get('career_objective', '').strip(),
            education=request.POST.get('education', '').strip(),
            experience=request.POST.get('experience', '').strip(),
            skills=request.POST.get('skills', '').strip(),
            projects=request.POST.get('projects', '').strip(),
            certifications=request.POST.get('certifications', '').strip(),
            languages=request.POST.get('languages', '').strip(),
            achievements=request.POST.get('achievements', '').strip(),
            job_description=request.POST.get('job_description', '').strip(),
        )

        # Handle cropped photo (base64 data URL)
        cropped_data = request.POST.get('cropped_photo_data', '')
        if cropped_data and cropped_data.startswith('data:image'):
            try:
                fmt, imgstr = cropped_data.split(';base64,')
                ext = fmt.split('/')[-1]
                filename = f"photo_{uuid.uuid4().hex[:8]}.{ext}"
                resume.photo = ContentFile(base64.b64decode(imgstr), name=filename)
            except Exception:
                pass  # Skip photo if anything goes wrong

        resume.save()
        refresh_resume_ats_score(resume)
        return redirect('view_resume', pk=resume.pk)

    return render(request, 'resumes/create_resume.html', {
        'selected_template': ATS_TEMPLATE,
    })


# ✅ LIST VIEW
class ResumeListView(LoginRequiredMixin, ListView):
    model = Resume
    template_name = "resumes/my_resumes.html"

    def get_queryset(self):
        return Resume.objects.filter(user=self.request.user)


# ✅ UPDATE VIEW
class ResumeUpdateView(LoginRequiredMixin, UpdateView):
    model = Resume
    form_class = ResumeForm
    template_name = "resumes/create_resume.html"
    success_url = reverse_lazy('my_resumes')

    def get_queryset(self):
        return Resume.objects.filter(user=self.request.user)

    def form_valid(self, form):
        form.instance.template = ATS_TEMPLATE
        response = super().form_valid(form)
        refresh_resume_ats_score(self.object)
        return response

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        o = self.object
        ctx['selected_template'] = ATS_TEMPLATE
        ctx['edit_initial_json'] = json.dumps({
            'full_name': o.full_name,
            'job_title': o.job_title or '',
            'email': o.email,
            'phone': o.phone,
            'address': o.address or '',
            'linkedin': o.linkedin or '',
            'github': o.github or '',
            'portfolio': o.portfolio or '',
            'twitter': o.twitter or '',
            'career_objective': o.career_objective or '',
            'skills': o.skills or '',
            'projects': o.projects or '',
            'certifications': o.certifications or '',
            'languages': o.languages or '',
            'achievements': o.achievements or '',
            'job_description': o.job_description or '',
            'education_raw': o.education or '',
            'experience_raw': o.experience or '',
            'photo_url': o.photo.url if o.photo else '',
        })
        return ctx


# ✅ DELETE VIEW
class ResumeDeleteView(LoginRequiredMixin, DeleteView):
    model = Resume
    template_name = "resumes/delete_resume.html"
    success_url = reverse_lazy('my_resumes')

    def get_queryset(self):
        return Resume.objects.filter(user=self.request.user)


@login_required
@require_http_methods(["POST"])
def delete_resume_ajax(request, pk):
    """
    Delete a resume from the list page (modal confirmation).
    Returns JSON for modern UX; avoids navigating to a confirmation page.
    """
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    resume.delete()
    return JsonResponse({"ok": True})

# ✅ DETAIL VIEW
class ResumeDetailView(LoginRequiredMixin, DetailView):
    model = Resume

    def get_queryset(self):
        return Resume.objects.filter(user=self.request.user)

    def get_template_names(self):
        return ["resumes/minimalist.html"]


# ✅ PDF GENERATION
def _pdf_link_callback(uri, rel):
    """
    Allow xhtml2pdf to resolve local media/static paths.
    """
    if uri.startswith(settings.MEDIA_URL):
        path = os.path.join(settings.MEDIA_ROOT, uri.replace(settings.MEDIA_URL, ""))
        return path
    if uri.startswith(settings.STATIC_URL):
        for static_dir in getattr(settings, "STATICFILES_DIRS", []):
            candidate = os.path.join(str(static_dir), uri.replace(settings.STATIC_URL, ""))
            if os.path.exists(candidate):
                return candidate
        return os.path.join(str(settings.BASE_DIR), "static", uri.replace(settings.STATIC_URL, ""))
    return uri


@login_required
def download_resume_pdf(request, pk):
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    template = get_template("resumes/pdf/minimalist.html")
    
    # Handle photo path for PDF generation - convert to absolute file path
    photo_path = None
    if resume.photo and resume.photo.name:
        photo_path = os.path.join(settings.MEDIA_ROOT, resume.photo.name)
    
    html = template.render({
        "resume": resume,
        "photo_path": photo_path,  # Pass the file path instead of URL
        "media_root": settings.MEDIA_ROOT,
        "static_root": settings.STATICFILES_DIRS[0] if settings.STATICFILES_DIRS else settings.BASE_DIR / 'static',
    })

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="resume_{resume.id}.pdf"'
    pisa_status = pisa.CreatePDF(html, dest=response, link_callback=_pdf_link_callback)

    if pisa_status.err:
        return HttpResponse('Error generating PDF')
    return response


# ✅ DOCX GENERATION
@login_required
def download_resume_docx(request, pk):
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    try:
        from docx import Document
    except Exception as e:
        return HttpResponse(f"python-docx is not installed: {e}", status=500)

    doc = Document()
    doc.add_heading(resume.full_name or "Resume", level=0)
    contact = " | ".join([s for s in [resume.email, resume.phone, resume.address] if s])
    if contact:
        doc.add_paragraph(contact)

    if resume.career_objective:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(resume.career_objective)

    if resume.skills:
        doc.add_heading("Skills", level=1)
        skills = [s.strip() for s in resume.skills.split(",") if s.strip()]
        if skills:
            for s in skills:
                doc.add_paragraph(s, style="List Bullet")
        else:
            doc.add_paragraph(resume.skills)

    exp_entries = resume.exp_entries() if hasattr(resume, "exp_entries") else []
    if resume.experience:
        doc.add_heading("Experience", level=1)
        if exp_entries:
            for e in exp_entries:
                head = " — ".join([x for x in [e.get("title"), e.get("company")] if x])
                dates = " – ".join([x for x in [e.get("start"), e.get("end")] if x])
                p = doc.add_paragraph(head)
                if dates:
                    p.add_run(f" ({dates})")
                if e.get("desc"):
                    doc.add_paragraph(e["desc"])
        else:
            doc.add_paragraph(resume.experience)

    edu_entries = resume.edu_entries() if hasattr(resume, "edu_entries") else []
    if resume.education:
        doc.add_heading("Education", level=1)
        if edu_entries:
            for e in edu_entries:
                head = " — ".join([x for x in [e.get("degree"), e.get("institution")] if x])
                dates = " – ".join([x for x in [e.get("start"), e.get("end")] if x])
                doc.add_paragraph(" ".join([head, f"({dates})" if dates else ""]).strip())
        else:
            doc.add_paragraph(resume.education)

    if resume.projects:
        doc.add_heading("Projects", level=1)
        doc.add_paragraph(resume.projects)
    if resume.certifications:
        doc.add_heading("Certifications", level=1)
        doc.add_paragraph(resume.certifications)
    if resume.achievements:
        doc.add_heading("Achievements", level=1)
        doc.add_paragraph(resume.achievements)
    if resume.languages:
        doc.add_heading("Languages", level=1)
        doc.add_paragraph(resume.languages)

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="resume_{resume.id}.docx"'
    doc.save(response)
    return response


# ✅ ATS ANALYZER
@login_required
def ats_analyzer(request, pk):
    resume = get_object_or_404(Resume, pk=pk, user=request.user)

    if request.method == "POST":
        job_description = request.POST.get("job_description", "").strip()
        resume.job_description = job_description
        resume.save(update_fields=["job_description"])
        refresh_resume_ats_score(resume)

    return render(request, "resumes/ats.html", {"resume": resume})


# ✅ PORTFOLIO VIEW (by pk)
@login_required
def portfolio_view(request, pk):
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    public_url = request.build_absolute_uri(
        reverse("public_portfolio", kwargs={"slug": resume.slug})
    )
    return render(
        request,
        "resumes/portfolio.html",
        {"resume": resume, "public_url": public_url, "is_public": False},
    )


# ✅ PUBLIC PORTFOLIO (by slug) — share this URL with anyone; no login required
def public_portfolio(request, slug):
    resume = get_object_or_404(Resume, slug=slug)
    return render(
        request,
        "resumes/portfolio.html",
        {"resume": resume, "is_public": True},
    )


# ✅ DASHBOARD
@login_required
def dashboard(request):
    resumes = Resume.objects.filter(user=request.user)
    return render(request, "resumes/dashboard.html", {"resumes": resumes})


# ✅ AI SUGGESTIONS
@login_required
def ai_suggestions(request, pk):
    resume = get_object_or_404(Resume, id=pk, user=request.user)
    prompt = build_improvement_prompt(
        career_objective=resume.career_objective,
        skills=resume.skills,
        experience=resume.experience,
    )
    result = generate_resume(prompt)
    return render(request, "resumes/ai_suggestions.html", {
        "resume": resume,
        "suggestion": result,
    })


# ✅ REAL-TIME ATS SCORE API (AI)
@login_required
@require_http_methods(["POST"])
def ats_score_api(request):
    """
    Accepts draft resume fields + optional job description and returns:
    { score: int(0-100), suggestions: [..], keywords: [..] }
    """
    try:
        payload = json.loads(request.body.decode("utf-8") or "{}")
    except Exception:
        payload = {}

    resume_text = payload.get("resume_text", "") or ""
    job_description = payload.get("job_description", "") or ""

    prompt = build_ats_prompt(resume_text=resume_text, job_description=job_description)
    raw = generate_resume(prompt)

    # Best-effort JSON parsing; fall back to heuristics if model returns plain text
    score = 0
    suggestions = []
    keywords = []
    try:
        data = json.loads(raw)
        score = int(data.get("score", 0))
        suggestions = data.get("suggestions", []) or []
        keywords = data.get("missing_keywords", []) or []
    except Exception:
        score = heuristic_ats_score(resume_text, job_description)
        suggestions = [raw] if raw else []

    score = max(0, min(score, 100))
    return JsonResponse({"score": score, "suggestions": suggestions, "missing_keywords": keywords})


def extract_uploaded_resume_text(uploaded):
    """Extract plain text from .txt, .pdf, or .docx uploads."""
    name = uploaded.name.lower()
    raw = uploaded.read()
    if name.endswith(".txt"):
        return raw.decode("utf-8", errors="replace")
    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(raw))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception:
            return ""
    if name.endswith(".docx"):
        try:
            from docx import Document

            doc = Document(BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""
    return ""


@login_required
def upload_resume_ats(request):
    """
    Upload an existing resume file and get an ATS score without saving a builder resume.
    """
    result_score = None
    error_message = None
    snippet = ""
    jd = ""

    if request.method == "POST":
        jd = request.POST.get("job_description", "").strip()
        f = request.FILES.get("resume_file")
        if not f:
            error_message = "Please upload a file (.pdf, .txt, or .docx)."
        elif f.size > 2 * 1024 * 1024:
            error_message = "File is too large (maximum 2 MB)."
        else:
            text = extract_uploaded_resume_text(f)
            if len((text or "").strip()) < 25:
                error_message = (
                    "Could not read enough text from this file. "
                    "Try exporting as .txt or use a text-based (not scanned-only) PDF."
                )
            else:
                result_score = compute_ats_score(text, jd)
                snippet = text.strip()[:3500]

    return render(
        request,
        "resumes/upload_resume_ats.html",
        {
            "result_score": result_score,
            "error_message": error_message,
            "snippet": snippet,
            "job_description": jd,
        },
    )


# ✅ TEMPLATE LIST / LANDING
def template_list(request):
    return render(request, "landing.html", {"resume_type": "ATS-perfect"})
